# ─────────────────────────────────────────────────────────────────────────────
# CSV to Document Converter
# ─────────────────────────────────────────────────────────────────────────────
# This script allows users to upload CSV files and convert them into different document formats (CSV, HTML, DOCX).
# It also provides options for grouping, sorting, and filtering the data before conversion.
# In the process of developing a site for documentation, via GitHub Pages,
# I wanted to create a simple and user-friendly interface for converting CSV files into various document formats.
# So developed this script to help users easily manage and convert their CSV data into more readable formats.
# It uses Streamlit for the web interface and Pandas for data manipulation.

import traceback
import streamlit as st
import pandas as pd
import io
import base64
import json
import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# ---------------------------------------------------------------------------
# Image handling utilities
# ---------------------------------------------------------------------------

BASE_DIR   = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR / "assets"

def get_float_image_path() -> str:
    """
    Returns an absolute path to “float_image.png” located inside ./assets.
    If the file does not exist, a harmless online placeholder is returned so
    that the UI never shows a broken-image icon.
    """
    img_path = ASSETS_DIR / "float_image.png"

    if img_path.exists():
        return str(img_path)

    # Fallback (any publicly available placeholder will do)
    return "https://via.placeholder.com/150"


# Replace the former hard-coded value
float_image_path: str = get_float_image_path()

# Set page config - must be the first Streamlit command
st.set_page_config(page_title="CSV to Document Converter", layout="wide", page_icon="favicon.png")

float_image_path = os.path.join(os.path.dirname(__file__), "assets", "favicon_transparent.png")
st.write(f"Image path: {float_image_path}")
# st.image(float_image_path, use_container_width=False, caption="Your Image", width=150)

# Create a data directory if it doesn't exist
def ensure_data_dir_exists():
    os.makedirs("data", exist_ok=True)
    os.makedirs("data/users", exist_ok=True)
    os.makedirs("data/csv_cache", exist_ok=True)

# Ensure the data directory exists
ensure_data_dir_exists()

# User management functions
def get_user_id():
    """Get a unique user ID or create one if it doesn't exist"""
    # Check if user_id exists in session state
    if "user_id" not in st.session_state:
        # Try to get user_id from query parameters (cookie-like behavior)
        if "user_id" in st.query_params and st.query_params["user_id"]:
            # Use the user_id from query parameters
            st.session_state.user_id = st.query_params["user_id"][0]
        elif "persistent_user_id" in st.session_state:
            # Use previously stored ID
            st.session_state.user_id = st.session_state.persistent_user_id
        else:
            # Generate a new user ID
            st.session_state.user_id = str(uuid.uuid4())
            st.session_state.persistent_user_id = st.session_state.user_id
            # Store user_id in query parameters (cookie-like behavior)
            st.query_params["user_id"] = st.session_state.user_id

    # Ensure user_id is always in query parameters
    if "user_id" not in st.query_params or not st.query_params["user_id"]:
        st.query_params["user_id"] = st.session_state.user_id

    return st.session_state.user_id

def get_user_history_path(user_id):
    """Get path to user's history file"""
    return f"data/users/{user_id}/history.json"

def save_to_user_history(user_id, file_name, cache_id=None):
    """
    Save file information to user history
    If cache_id is provided, it will be used to link to the cached file
    """
    # Choose an appropriate path for user history
    history_path = get_user_history_path(user_id)
    os.makedirs(os.path.dirname(history_path), exist_ok=True)
    user_history = []

    # Load existing history if it exists
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            user_history = json.load(f)

    # Check if this file is already in history
    file_exists = False
    for entry in user_history:
        if isinstance(entry, dict) and entry.get("file_name") == file_name:
            # Update the existing entry with new timestamp and cache_id
            entry["timestamp"] = datetime.now().isoformat()
            if cache_id:
                entry["cache_id"] = cache_id
            file_exists = True
            break

    # If file doesn't exist in history, add it
    if not file_exists:
        csv_data = {
            "file_name": file_name,
            "timestamp": datetime.now().isoformat()
        }
        if cache_id:
            csv_data["cache_id"] = cache_id
        user_history.append(csv_data)

    # Save updated history
    with open(history_path, 'w', encoding='utf-8') as f:
        json.dump(user_history, f)

def get_user_history(user_id):
    history_path = get_user_history_path(user_id)
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

# CSV cache management
def ensure_csv_cache_dir_exists():
    """Create the CSV cache directory if it doesn't exist"""
    os.makedirs("data/csv_cache", exist_ok=True)

def get_file_hash(df):
    """Generate a hash for the DataFrame content to use as identifier"""
    # Convert DataFrame to CSV string and hash it
    csv_string = df.to_csv(index=False)
    return str(hash(csv_string))

def cache_csv(df, file_name):
    """
    Save DataFrame to session cache and disk cache with its file name
    Returns a cache_id that can be used to retrieve the DataFrame later
    """
    # Ensure cache directory exists
    ensure_csv_cache_dir_exists()

    # Generate a hash for the file content
    file_hash = get_file_hash(df)

    # Check if we already have this file in the cache
    cache_path = f"data/csv_cache/{file_hash}.csv"

    # Save to disk cache if not already there
    if not os.path.exists(cache_path):
        df.to_csv(cache_path, index=False)

    # Update the uploaded_files_history.json
    history_path = "uploaded_files_history.json"
    history = {}
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                history = json.load(f)
            except json.JSONDecodeError:
                history = {}

    # Store file info in history
    history[file_hash] = {
        "file_name": file_name,
        "cache_path": cache_path,
        "timestamp": datetime.now().isoformat()
    }

    with open(history_path, 'w', encoding='utf-8') as f:
        json.dump(history, f)

    # Also store in session cache for current session
    if "csv_cache" not in st.session_state:
        st.session_state.csv_cache = {}

    st.session_state.csv_cache[file_hash] = {
        "df": df,
        "file_name": file_name,
        "timestamp": datetime.now().isoformat()
    }

    return file_hash

def get_cached_csv(cache_id):
    """
    Retrieve DataFrame from cache (session or disk)
    Returns the DataFrame or None if not found
    """
    # First try session cache
    if "csv_cache" in st.session_state and cache_id in st.session_state.csv_cache:
        return st.session_state.csv_cache[cache_id]["df"]

    # If not in session, try disk cache
    cache_path = f"data/csv_cache/{cache_id}.csv"
    if os.path.exists(cache_path):
        try:
            df = pd.read_csv(cache_path, dtype=str)

            # Add to session cache for future use
            if "csv_cache" not in st.session_state:
                st.session_state.csv_cache = {}

            # Get file name from history
            file_name = get_file_name_from_history(cache_id)

            st.session_state.csv_cache[cache_id] = {
                "df": df,
                "file_name": file_name,
                "timestamp": datetime.now().isoformat()
            }

            return df
        except Exception as e:
            st.error(f"Error loading cached CSV: {e}")

    return None

def get_file_name_from_history(cache_id):
    """Get the original file name from the history"""
    history_path = "uploaded_files_history.json"
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                history = json.load(f)
                if cache_id in history:
                    return history[cache_id]["file_name"]
            except json.JSONDecodeError:
                pass

    return f"cached_file_{cache_id}.csv"

def get_all_cached_csvs():
    """Get all cached CSVs for the current session"""
    if "csv_cache" not in st.session_state:
        st.session_state.csv_cache = {}
    return st.session_state.csv_cache

def get_all_disk_cached_csvs():
    """Get all CSVs from the disk cache"""
    history_path = "uploaded_files_history.json"
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                pass

    return {}

# Your existing conversion functions
def convert_df_to_csv(df):
    """Convert DataFrame to CSV string"""
    return df.to_csv(index=False).encode('utf-8')

# python
def convert_df_to_html(df, doc_title=None):
    """Convert DataFrame to HTML string with an optional title"""
    html_table = df.to_html(index=False, border=0)
    style = """
    <style>
      table, th, td {
        border: 1px solid black;
        border-collapse: collapse;
        padding: 5px 10px;
      }
    </style>
    """
    title_html = f"<h1>{doc_title}</h1>" if doc_title else ""
    return style + title_html + html_table

def convert_df_to_grouped_html(df: pd.DataFrame, group_cols: list[str], doc_title: str) -> str:
    """
    Return an HTML string in which *df* is split by the given columns.
    If *group_cols* is empty, the plain table is returned.
    """
    if not group_cols:
        return convert_df_to_html(df)

    html_chunks: list[str] = []
    grouped = df.groupby(group_cols)

    if doc_title:
        title = f"<h3>{doc_title}</h3>"
    else:
        title = ", ".join(f"{col} = {val}" for col, val in group_cols.items())
    previous_group = None
    for keys, group in grouped:
        if previous_group is not None and previous_group != group:
            previous_group = group
            # Add a blank line between groups
            html_chunks.append("<br>")
            html_chunks.append(f"<h3>{title}</h3>")
        # Normalise keys → always a tuple so we can zip()
        keys = (keys,) if not isinstance(keys, tuple) else keys
        html_chunks.append(group.to_html(index=False))

    return "\n".join(html_chunks)


def convert_df_to_docx(df, title="Data Document"):
    """Convert DataFrame to DOCX"""
    doc = Document()
    doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add table
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        run = hdr_cells[i].paragraphs[0].add_run(str(col))
        run.bold = True
        run.font.size = Pt(12)

    # Data rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            run = cells[i].paragraphs[0].add_run(str(row[col]))
            run.font.size = Pt(10)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def convert_df_to_grouped_docx(df: pd.DataFrame, group_cols: list[str], doc_title:str) -> bytes:
    """
    Build a DOCX file where *df* is split by *group_cols*.
    Returns the binary content (bytes).
    """
    if not group_cols:
        return convert_df_to_docx(df)

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    grouped = df.groupby(group_cols)

    for keys, group in grouped:
        keys = (keys,) if not isinstance(keys, tuple) else keys
        if doc_title:
            title = doc_title
        else:
            heading = ", ".join(f"{col} = {val}" for col, val in zip(group_cols, keys))
        doc.add_heading(heading, level=3)

        # Add the table for this slice
        table = doc.add_table(rows=1, cols=len(group.columns))
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        for idx, col in enumerate(group.columns):
            hdr_cells[idx].text = str(col)
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for _, row in group.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.add_paragraph()  # blank line between groups

    # Serialize to bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def get_download_link(data, filename, text):
    """Generate a download link for the data"""
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:file/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

# Assume: df is your current DataFrame
# download_format: "csv", "html", or "docx"

def get_data_preview_and_download_options(df, download_format):
    # Generate preview HTML based on the format
    if download_format == "csv":
        preview = df.head().to_csv()
        file_content = convert_df_to_csv(df)
        file_ext = "csv"
    elif download_format == "html":
        preview = df.head().to_html()
        file_content = convert_df_to_html(df)
        file_ext = "html"
    elif download_format == "docx":
        preview = "Preview not available for DOCX format." # Or show a summary
        file_content = convert_df_to_docx(df)
        file_ext = "docx"
    else:
        preview = ""
        file_content = None
        file_ext = ""

    # Always return a single preview + download section
    download_link = get_download_link(file_content, f"export.{file_ext}", f"Download File - Table")
    return f"""
        <div id='data-preview'>
            <h4>Preview:</h4>
            <pre>{preview}</pre>
            <h4>Download:</h4>
            {download_link}
        </div>
    """

# In your route or callback or rendering logic:
# Then, in the template/UI rendering, always inject/replace the data_preview_section

# ... keep your existing conversion functions unchanged ...

# Initialize session state for storing dataframes if not exists
if "dataframes" not in st.session_state:
    st.session_state.dataframes = []

# Get the current user ID
user_id = get_user_id()

col_img, col_body = st.columns([1, 10])      # adjust ratios to your liking

with col_img:
    st.image(float_image_path, width=150)

with col_body:
    st.write("""
        This text appears to the right of the image, so it feels very
        similar to an HTML float-left.  
        You can put any other Streamlit widgets here as well.
    """)

    st.title("CSV to Document Converter")
    st.write("""
    Upload one or more CSV files. Each loaded CSV is displayed in its own table.
    You can convert each to different document formats.
    """)

    # Display cached CSVs for the current session
    with st.expander("Current Session Cache", expanded=False):
        cached_csvs = get_all_cached_csvs()
        if cached_csvs:
            st.write("CSVs cached in current session:")
            for cache_id, cache_info in cached_csvs.items():
                # Create a button to reload this CSV
                if st.button(f"Load: {cache_info['file_name']}", key=f"load_{cache_id}"):
                    # Add this cached CSV to dataframes for processing
                    df = cache_info['df']
                    if df is not None and not any(df.equals(existing_df) for existing_df in st.session_state.dataframes):
                        st.session_state.dataframes.append(df)
                        st.success(f"Loaded cached CSV: {cache_info['file_name']}")
                        st.rerun()
        else:
            st.write("No CSVs are cached in the current session.")

    # File uploader
    uploaded_file = st.file_uploader("Choose a CSV file", type="csv", key="upload")

    if uploaded_file is not None:
        try:
            df = pd.read_csv(uploaded_file, dtype=str)

            # Add to dataframes list
            st.session_state.dataframes.append(df)

            # Cache the CSV and get the cache_id
            cache_id = cache_csv(df, uploaded_file.name)

            # Add to user history with the cache_id
            save_to_user_history(user_id, uploaded_file.name, cache_id)

            st.success(f"Successfully loaded CSV with {df.shape[0]} rows and {df.shape[1]} columns.")
            # download_format = st.selectbox("Format", ["csv", "html", "docx"])
            # enable_grouping = st.checkbox("Group by column")
            # if enable_grouping:
            #     # switch to a single‐select so df[group_col] is a Series
            #     group_col = st.selectbox("Select column to group by", df.columns)
            #     unique_vals = df[group_col].unique().tolist()
            #     # let user pick exactly one group
            #     chosen = st.selectbox("Which group to preview?", unique_vals)
            #     subset = df[df[group_col] == chosen]
            #     st.markdown(
            #         get_data_preview_and_download_options(subset, download_format),
            #         unsafe_allow_html=True
            #     )
            # else:
            #     st.markdown(
            #         get_data_preview_and_download_options(df, download_format),
            #         unsafe_allow_html=True
            #     )

        except Exception as e:
            st.error(f"Error processing the file")
            st.error(f"Exception occurred")
            st.markdown(f"**Exception:**<br>{traceback.format_exc().replace('\n', '<br>')}", unsafe_allow_html=True)

    # Display all loaded CSVs in separate tables
    for idx, df in enumerate(st.session_state.dataframes):
        with st.expander(f"Data Table", expanded=True):
            st.subheader(f"Data Preview - Table")
            st.text(f"Loaded {df.shape[0]} rows and {df.shape[1]} columns.")
            st.dataframe(df.head(10))

            # Document title input per dataframe
            doc_title = st.text_input(f"Document Title for Table",
                                     f"Data Document {idx+1}",
                                     key=f"title_{idx}")

            # Grouping, sorting, and filtering in tabs
            tab1, tab2, tab3 = st.tabs(["Grouping", "Sorting", "Filtering"])

            with tab1:
                enable_grouping = st.checkbox("Enable grouping by column", key=f"grouping_{idx}")
                group_cols = []
                if enable_grouping and len(df.columns) > 0:
                    group_cols = st.multiselect("Select column(s) to group by", df.columns, key=f"group_col_{idx}")

            with tab2:
                enable_sorting = st.checkbox("Enable sorting", key=f"sort_enable_{idx}")
                sort_cols = []
                ascending = True
                if enable_sorting:
                    sort_cols = st.multiselect("Select column(s) to sort by",
                                            options=df.columns,
                                            key=f"sort_cols_{idx}")
                    sort_order = st.selectbox("Sort order", ("Ascending", "Descending"), key=f"sort_order_{idx}")
                    ascending = sort_order == "Ascending"

            with tab3:
                enable_filtering = st.checkbox("Enable filtering", key=f"filter_enable_{idx}")
                filter_conditions = {}
                include_columns = []
                exclude_columns = []

                if enable_filtering:
                    # Value filtering
                    st.subheader("Filter by values")
                    filter_cols = st.multiselect("Select column(s) to filter by value",
                                              options=df.columns,
                                              key=f"filter_cols_{idx}")
                    for col in filter_cols:
                        unique_values = df[col].dropna().unique().tolist()
                        selected_values = st.multiselect(f"Select values for {col}",
                                                     options=unique_values,
                                                     key=f"filter_values_{idx}_{col}")
                        if selected_values:
                            filter_conditions[col] = selected_values

                    # Column filtering
                    st.subheader("Filter by columns")
                    column_filter_type = st.radio(
                        "Column filtering type",
                        ["None", "Include only", "Exclude only"],
                        key=f"column_filter_type_{idx}"
                    )

                    if column_filter_type == "Include only":
                        include_columns = st.multiselect(
                            "Select columns to include in the output",
                            options=df.columns,
                            key=f"include_columns_{idx}"
                        )
                    elif column_filter_type == "Exclude only":
                        exclude_columns = st.multiselect(
                            "Select columns to exclude from the output",
                            options=df.columns,
                            key=f"exclude_columns_{idx}"
                        )

            # Apply transformations to create the processed dataframe
            processed_df = df.copy()

            # Apply value filtering
            for col, values in filter_conditions.items():
                processed_df = processed_df[processed_df[col].isin(values)]

            # Apply column filtering (include or exclude)
            if include_columns:
                processed_df = processed_df[include_columns]
            elif exclude_columns:
                processed_df = processed_df.drop(columns=exclude_columns)

            # Apply sorting
            if enable_sorting and sort_cols:
                processed_df = processed_df.sort_values(by=sort_cols, ascending=ascending)

            # Download options
            st.subheader(f"Download Options for Table")
            download_format = st.radio("Select download format",
                                     ["CSV", "HTML", "DOCX"],
                                     key=f"download_format_{idx}")

            if st.button(f"Generate Document for Table"):
                if download_format == "CSV":
                    csv_data = convert_df_to_csv(processed_df)
                    st.download_button(
                        label=f"Download CSV - Table",
                        data=csv_data,
                        file_name=f"data_export_{idx+1}.csv",
                        mime="text/csv",
                        on_click="ignore"
                    )
                elif download_format == "HTML":
                    if enable_grouping and group_cols:
                        html_data = convert_df_to_grouped_html(processed_df, group_cols, doc_title)
                        html_bytes = html_data.encode('utf-8')  # Convert string to bytes
                    else:
                        html_data = convert_df_to_html(processed_df, doc_title)
                        html_bytes = html_data.encode('utf-8') if isinstance(html_data, str) else html_data

                    st.download_button(
                        label=f"Download HTML - Table",
                        data=html_bytes,
                        file_name=f"grouped_data_{idx+1}.html" if enable_grouping and group_cols else f"data_{idx+1}.html",
                        mime="text/html",
                        on_click="ignore"
                    )
                elif download_format == "DOCX":
                    # Similar handling for DOCX format
                    if enable_grouping and group_cols:
                        docx_data = convert_df_to_grouped_docx(processed_df, group_cols, doc_title)
                    else:
                        docx_data = convert_df_to_docx(processed_df, doc_title)
                    st.download_button(
                        label=f"Download DOCX - Table",
                        data=docx_data,
                        file_name=f"grouped_data_{idx+1}.docx" if enable_grouping and group_cols else f"data_{idx+1}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        on_click="ignore"
                    )
        break

    # Footer
    st.markdown("---")
    st.markdown("CSV to Document Converter | Created with Streamlit")
    st.markdown("This app allows you to upload a CSV file and convert it to different document formats. You can also group the data by a column to create separate tables for each unique value.")
    st.markdown("**Note:** Make sure to upload a valid CSV file.")
    st.markdown("**Disclaimer:** This app is for educational purposes only. Please ensure you have the right to use any data you upload.")
    # ─────────────────────────────────────────────────────────────────────────────
    #  CONVERT A DATAFRAME TO GROUPED HTML
    # ─────────────────────────────────────────────────────────────────────────────
    def convert_df_to_grouped_html(
        df: pd.DataFrame,
        group_cols: list[str],
        doc_title: str | None = None,        # ← NEW, optional third parameter
    ) -> str:
        """
        Convert a DataFrame to an HTML string, grouped by the specified columns.

        Parameters
        ----------
        df         : pd.DataFrame
            The data that will be rendered.
        group_cols : list[str]
            Column names to group by.
        doc_title  : str | None, optional
            An optional document title.  If supplied, a corresponding
            <h1> header is inserted at the top of the generated HTML.

        Returns
        -------
        str
            The rendered HTML string.
        """
        html_parts: list[str] = []

        # Add an overall title when requested
        if doc_title:
            html_parts.append(f"<h1>{doc_title}</h1>")

        # Group the DataFrame and render each subgroup
        grouped = df.groupby(group_cols, dropna=False)

        for keys, group in grouped:
            # Ensure *keys* is always iterable
            keys = (keys,) if not isinstance(keys, tuple) else keys
            group_name = ", ".join(f"{col}: {val}" for col, val in zip(group_cols, keys))

            html_parts.append(f"<h2>{group_name}</h2>")
            html_parts.append(
                group.to_html(index=False, escape=False, border=0, classes="dataframe")
            )

        return "\n".join(html_parts)

    # Display user history in an expander
    with st.expander("Your CSV History", expanded=False):
        try:
            user_history = get_user_history(user_id)
            if user_history:
                st.write("Previously uploaded CSV files (click to load):")
                for idx, entry in enumerate(user_history):
                    if isinstance(entry, dict) and 'file_name' in entry and 'timestamp' in entry:
                        # If entry has cache_id, make it clickable
                        if 'cache_id' in entry:
                            if st.button(f"Load: {entry['file_name']}", key=f"history_{idx}"):
                                # Load the cached CSV
                                df = get_cached_csv(entry['cache_id'])
                                if df is not None:
                                    # Add to dataframes list if not already there
                                    if not any(df.equals(existing_df) for existing_df in st.session_state.dataframes):
                                        st.session_state.dataframes.append(df)
                                        st.success(f"Loaded {entry['file_name']} from history")
                                        st.rerun()
                                    else:
                                        st.info(f"{entry['file_name']} is already loaded")
                                else:
                                    st.error(f"Could not load {entry['file_name']} from cache")
                            # Display timestamp separately
                            st.caption(f"Uploaded: {entry['timestamp']}")
                        else:
                            # For entries without cache_id, just display them
                            st.write(f"{idx+1}. **{entry['file_name']}** - {entry['timestamp']}")
                    elif isinstance(entry, str):
                        # Handle legacy entries that might only contain filenames
                        st.write(f"{idx+1}. **{entry}**")
                    else:
                        st.write(f"{idx+1}. **Unknown entry format**")
            else:
                st.write("No CSV files have been uploaded yet.")
        except Exception as e:
            st.error(f"Error loading user history: {e}")
            # Option 1: Just show plain error (preferred)
            st.error(f"Exception: {traceback.format_exc()}")
            # Option 2: For HTML formatting, use markdown (if you really need HTML line breaks)
            # st.markdown(f"**Exception:**<br>{traceback.format_exc().replace(chr(10), '<br>')}", unsafe_allow_html=True)
