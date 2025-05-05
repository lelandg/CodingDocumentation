import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load the CSV you already have
df = pd.read_csv('emojis.csv', dtype=str)

doc = Document()
doc.add_heading("Emoji Alt Key Combinations", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

# For each emoji group, make a heading and a table
for group_name, group_df in df.groupby('Group'):
    doc.add_heading(group_name, level=2)
    cols = list(group_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        run = hdr_cells[i].paragraphs[0].add_run(col)
        run.bold = True
        run.font.size = Pt(15)
    # Data rows
    for _, row in group_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            run = cells[i].paragraphs[0].add_run(str(row[col]))
            run.font.size = Pt(15)
    doc.add_paragraph()  # spacing

doc.save('emojis_grouped.docx')
print("Saved emojis_grouped.docx")
